import cv2
import matplotlib.pyplot as plt
from tensorflow.keras.preprocessing.image import img_to_array
from sklearn.preprocessing import LabelEncoder
from sklearn.metrics import classification_report
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import Conv2D, MaxPooling2D, Flatten, Dense
from docx import Document

#  Image Loading and Preprocessing
# Load a sample sad human image
image_path_sad = "image.jpg"
image_sad = cv2.imread(image_path_sad)
image_rgb_sad = cv2.cvtColor(image_sad, cv2.COLOR_BGR2RGB)

# Display the sample sad image
plt.imshow(image_rgb_sad)
plt.axis('off')
plt.show()

# Preprocess the sad image
target_size = (224, 224)
resized_image_sad = cv2.resize(image_rgb_sad, target_size)
resized_image_sad = img_to_array(resized_image_sad) / 255.0  # Normalize pixel values

#  Label Encoding
# Encode emotion label 'sad' into a numerical value
emotion_labels_sad = ['sad']
label_encoder_sad = LabelEncoder()
encoded_label_sad = label_encoder_sad.fit_transform(emotion_labels_sad)

# Reverse mapping to get the original label
original_label_sad = label_encoder_sad.inverse_transform(encoded_label_sad)


#  Model Selection
model_sad = Sequential()
model_sad.add(Conv2D(32, (3, 3), activation='relu', input_shape=(224, 224, 3)))
model_sad.add(MaxPooling2D((2, 2)))
model_sad.add(Conv2D(64, (3, 3), activation='relu'))
model_sad.add(MaxPooling2D((2, 2)))
model_sad.add(Flatten())
model_sad.add(Dense(64, activation='relu'))
model_sad.add(Dense(1, activation='sigmoid'))  # Binary classification for 'sad'

#  Model Compilation
model_sad.compile(optimizer='adam', loss='binary_crossentropy', metrics=['accuracy'])

#  Model Training
model_sad.fit(resized_image_sad.reshape(1, 224, 224, 3), encoded_label_sad, epochs=10, batch_size=1)



#  Model Evaluation
y_pred_sad = model_sad.predict(resized_image_sad.reshape(1, 224, 224, 3))
predicted_label_sad = label_encoder_sad.inverse_transform(y_pred_sad > 0.5)

# Create a Word document
doc = Document()
doc.add_heading('Model Evaluation Results', 0)

# Add evaluation metrics to the document
report = classification_report([encoded_label_sad], [int(y_pred_sad > 0.5)], output_dict=True)
for label, metrics in report.items():
    if label.isdigit():  # Exclude 'support' metrics
        doc.add_heading(f'Class: {label_encoder_sad.inverse_transform([int(label)])[0]}', level=1)
        for metric, value in metrics.items():
            doc.add_paragraph(f'{metric.capitalize()}: {value}')

# Save the document
doc.save('model_evaluation_results.docx')
